"""
Management command che genera i due template Word trilingui (IT/DE/EN) per i permessi.

    python manage.py create_permit_template

Genera:
  - permit_template_admin.docx   → copia gestore (cedolino/matrice)
  - permit_template_client.docx  → copia cliente (da esporre sul parabrezza)

I placeholder {{ variabile }} devono restare invariati.
"""

from django.core.management.base import BaseCommand
from django.conf import settings
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Costanti testo trilingue
# ---------------------------------------------------------------------------

HEADER = (
    "PROVINCIA AUTONOMA DI TRENTO\n"
    "ASUC DI CELENTINO\n"
    "Via Giovanni Caserotti n. 31 – 38024 Cogolo di Peio (TN)"
)

ROAD_NAME = "MALGA CAMPO"

# Titolo copia cliente
CLIENT_TITLE = [
    "AUTORIZZAZIONE AL TRANSITO SU STRADA FORESTALE DI TIPO B",
    "DURCHFAHRTSGENEHMIGUNG AUF FORSTSTRASSE TYP B",
    "TRANSIT AUTHORIZATION ON FOREST ROAD TYPE B",
]

# Nota legale
LEGAL_NOTE = [
    ("In base all'art. 100 della Legge Provinciale 23 maggio 2007, n. 11\n"
     "e relativo regolamento di attuazione, accertato che sussistono le condizioni\n"
     "di cui all'art. 29 del regolamento,"),
    ("Gemäß Art. 100 des Landesgesetzes vom 23. Mai 2007, Nr. 11\n"
     "und der entsprechenden Durchführungsverordnung, nachdem festgestellt wurde,\n"
     "dass die Voraussetzungen gemäß Art. 29 der Verordnung vorliegen,"),
    ("Pursuant to Art. 100 of Provincial Law of 23 May 2007, No. 11\n"
     "and its implementing regulation, having ascertained that the conditions\n"
     "set forth in Art. 29 of the regulation are met,"),
]

SI_AUTORIZZA = [
    "SI AUTORIZZA",
    "WIRD GENEHMIGT",
    "IS AUTHORIZED",
]

IL_TRANSITO = [
    "IL TRANSITO DEL VEICOLO:",
    "DIE DURCHFAHRT DES FAHRZEUGS:",
    "VEHICLE TRANSIT:",
]

SULLA_STRADA = [
    "sulla Strada Forestale denominata",
    "auf der Forststraße namens",
    "on the forest road named",
]

PER_IL_PERIODO = [
    "per il periodo: / für den Zeitraum: / for the period:",
]

DAL_AL_LABELS = ("dal / von / from", "al / bis / to")

FIRMA = [
    "Firma e/o timbro proprietario / titolare della gestione:",
    "Unterschrift und/oder Stempel des Eigentümers / Verwalters:",
    "Signature and/or stamp of the owner / manager:",
]

ESPORRE = [
    "⚠  DA ESPORRE IN MODO BEN VISIBILE SUL VEICOLO",
    "⚠  GUT SICHTBAR IM FAHRZEUG AUSZULEGEN",
    "⚠  TO BE DISPLAYED IN A CLEARLY VISIBLE POSITION IN THE VEHICLE",
]

# Titolo copia admin (cedolino)
ADMIN_TITLE = [
    "MATRICE RELATIVA ALL'AUTORIZZAZIONE AL TRANSITO",
    "MATRIX DER DURCHFAHRTSGENEHMIGUNG",
    "TRANSIT AUTHORIZATION MATRIX",
]
ADMIN_TITLE2 = [
    "SU STRADA FORESTALE DI TIPO B",
    "AUF FORSTSTRASSE TYP B",
    "ON FOREST ROAD TYPE B",
]

RILASCIA = [
    "Il proprietario/gestore della Strada Forestale di Tipo B denominata",
    "Der Eigentümer/Verwalter der Forststraße Typ B namens",
    "The owner/manager of the Type B Forest Road named",
]

AL_SIGNOR = [
    "rilascia al Signor/a:",
    "gewährt Herrn/Frau:",
    "grants to Mr./Ms.:",
]

IL_QUALE = [
    ("il quale ha dichiarato di dover percorrere la strada forestale sopra citata\n"
     "con il veicolo:"),
    ("der/die erklärt hat, die oben genannte Forststraße\n"
     "mit folgendem Fahrzeug befahren zu müssen:"),
    ("who has declared the need to travel on the aforementioned forest road\n"
     "with the following vehicle:"),
]

AUTH_PERIODO = [
    "L'AUTORIZZAZIONE AL TRANSITO PER IL PERIODO",
    "DIE DURCHFAHRTSGENEHMIGUNG FÜR DEN ZEITRAUM",
    "THE TRANSIT AUTHORIZATION FOR THE PERIOD",
]

# Etichette campi: IT / DE / EN
LABEL_NOME      = "Nome / Vorname / First name"
LABEL_COGNOME   = "Cognome / Nachname / Last name"
LABEL_RESIDENTE = "Residente a / Wohnhaft in / Resident in"
LABEL_VEICOLO   = "Tipo veicolo / Fahrzeugtyp / Vehicle type"
LABEL_TARGA     = "Targa / Kennzeichen / Plate"
LABEL_INTEST    = "Intestatario / Inhaber / Holder"
LABEL_DATA_IT   = "Data emissione / Ausstellungsdatum / Issue date"
LABEL_DATA_CL   = "Data / Datum / Date"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _new_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    style = doc.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(3)
    style.font.name = "Arial"
    style.font.size = Pt(11)
    return doc


def _sep(doc):
    p = doc.add_paragraph("─" * 60)
    p.runs[0].font.size = Pt(8)


def _heading(doc, text, size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)


def _body(doc, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p


def _trilingual_block(doc, lines, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.CENTER):
    """Aggiunge le tre traduzioni: IT in nero bold, DE e EN in grigio più piccolo."""
    _body(doc, lines[0], bold=True, size=size, align=align)
    if len(lines) > 1:
        _body(doc, lines[1], bold=False, size=size - 1, align=align, color=(100, 100, 100))
    if len(lines) > 2:
        _body(doc, lines[2], bold=False, size=size - 1, align=align, color=(130, 130, 130))


def _field(doc, label, placeholder, size=12):
    """Riga etichetta (piccola, grigia) + valore (bold)."""
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}:  ")
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(80, 80, 80)
    r2 = p.add_run(placeholder)
    r2.bold = True
    r2.font.size = Pt(size)
    return p


def _dal_al(doc, ph_dal, ph_al):
    """Riga 'dal / von / from  __  al / bis / to  __'."""
    p = doc.add_paragraph()
    r1 = p.add_run(f"{DAL_AL_LABELS[0]}  ")
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(80, 80, 80)
    r2 = p.add_run(ph_dal)
    r2.bold = True
    r2.font.size = Pt(12)
    p.add_run("      ")
    r3 = p.add_run(f"{DAL_AL_LABELS[1]}  ")
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(80, 80, 80)
    r4 = p.add_run(ph_al)
    r4.bold = True
    r4.font.size = Pt(12)


# ---------------------------------------------------------------------------
# Copia gestore – cedolino/matrice
# ---------------------------------------------------------------------------

def _build_admin_template(doc):
    _heading(doc, HEADER, size=10)
    doc.add_paragraph()

    _sep(doc)
    _trilingual_block(doc, ADMIN_TITLE, bold=True, size=13)
    _trilingual_block(doc, ADMIN_TITLE2, bold=True, size=12)
    _sep(doc)

    doc.add_paragraph()
    _trilingual_block(doc, RILASCIA, size=11, align=WD_ALIGN_PARAGRAPH.LEFT)
    _body(doc, ROAD_NAME, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    _trilingual_block(doc, AL_SIGNOR, size=11, align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()

    _field(doc, LABEL_NOME,      "{{ nome }}")
    _field(doc, LABEL_COGNOME,   "{{ cognome }}")
    _field(doc, LABEL_RESIDENTE, "{{ residenza }}")

    doc.add_paragraph()
    _trilingual_block(doc, IL_QUALE, size=11, align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()

    _field(doc, LABEL_VEICOLO, "{{ tipo_veicolo }}")
    _field(doc, LABEL_TARGA,   "{{ targa }}")

    doc.add_paragraph()
    _trilingual_block(doc, AUTH_PERIODO, bold=True, size=12)
    _dal_al(doc, "{{ dal }}", "{{ al }}")

    doc.add_paragraph()
    _sep(doc)

    p_num = doc.add_paragraph()
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_num.add_run("N.  {{ numero }} / {{ anno }}")
    r.bold = True
    r.font.size = Pt(18)

    _sep(doc)
    doc.add_paragraph()
    _field(doc, LABEL_DATA_IT, "{{ data_emissione }}")


# ---------------------------------------------------------------------------
# Copia cliente – autorizzazione parabrezza
# ---------------------------------------------------------------------------

def _build_client_template(doc):
    _heading(doc, HEADER, size=10)
    doc.add_paragraph()

    _sep(doc)
    _trilingual_block(doc, CLIENT_TITLE, bold=True, size=13)
    _sep(doc)

    doc.add_paragraph()
    for line in LEGAL_NOTE:
        _body(doc, line, size=10,
              color=(60, 60, 60) if LEGAL_NOTE.index(line) > 0 else None,
              align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()
    _trilingual_block(doc, SI_AUTORIZZA, bold=True, size=15)
    doc.add_paragraph()

    _field(doc, LABEL_INTEST,   "{{ nome }} {{ cognome }}")
    _field(doc, LABEL_RESIDENTE, "{{ residenza }}")

    doc.add_paragraph()
    _trilingual_block(doc, IL_TRANSITO, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.LEFT)

    _field(doc, LABEL_VEICOLO, "{{ tipo_veicolo }}")
    _field(doc, LABEL_TARGA,   "{{ targa }}")

    doc.add_paragraph()
    _trilingual_block(doc, SULLA_STRADA, size=11, align=WD_ALIGN_PARAGRAPH.LEFT)
    _body(doc, ROAD_NAME, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    _trilingual_block(doc, PER_IL_PERIODO, size=11, align=WD_ALIGN_PARAGRAPH.LEFT)
    _dal_al(doc, "{{ dal }}", "{{ al }}")

    doc.add_paragraph()
    for line in FIRMA:
        _body(doc, line, size=10,
              color=(60, 60, 60) if FIRMA.index(line) > 0 else None)
    doc.add_paragraph()
    _field(doc, LABEL_DATA_CL, "{{ data_emissione }}")

    doc.add_paragraph()
    _sep(doc)

    p_num = doc.add_paragraph()
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_num.add_run("N.  {{ numero }} / {{ anno }}")
    r.bold = True
    r.font.size = Pt(18)

    _sep(doc)
    doc.add_paragraph()
    for line in ESPORRE:
        _body(doc, line, bold=(ESPORRE.index(line) == 0), size=11,
              color=(60, 60, 60) if ESPORRE.index(line) > 0 else None,
              align=WD_ALIGN_PARAGRAPH.CENTER)


# ---------------------------------------------------------------------------
# Command
# ---------------------------------------------------------------------------

class Command(BaseCommand):
    help = "Genera i due template Word trilingui (IT/DE/EN) per i permessi"

    def handle(self, *args, **options):
        for attr, builder, label in [
            ("PERMIT_TEMPLATE_ADMIN_PATH",  _build_admin_template,  "admin (cedolino gestore)"),
            ("PERMIT_TEMPLATE_CLIENT_PATH", _build_client_template, "cliente (parabrezza)"),
        ]:
            path = getattr(settings, attr)
            path.parent.mkdir(parents=True, exist_ok=True)
            doc = _new_doc()
            builder(doc)
            doc.save(path)
            self.stdout.write(self.style.SUCCESS(f"✓ {label}  →  {path}"))

        self.stdout.write("\nI placeholder {{ variabile }} devono restare invariati.")
        self.stdout.write("Aprire in Word/LibreOffice per rifinire la grafica se necessario.")
